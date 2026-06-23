from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = REPO_ROOT.parent
REPORT_ROOT = PROJECT_ROOT / "报告文档" / "v2-最终高分版"
OUTPUT_ROOT = REPORT_ROOT / "正式文档"

DEFAULT_DOCS = [
    "01_软件需求说明报告.md",
    "02_产品功能设计与原型报告.md",
    "03_数据库设计文档.md",
    "04_后台接口设计文档.md",
    "05_软件实现说明.md",
    "06_软件使用手册.md",
    "07_迭代规划记录及项目总结.md",
    "08_黑盒白盒测试文档.md",
    "09_答辩PPT大纲与演示脚本.md",
]

OUTPUT_NAME_OVERRIDES = {
    "09_答辩PPT大纲与演示脚本.md": "09_答辩演示脚本.docx",
}

HEADING_COLOR = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK_COLOR = RGBColor(0x1F, 0x4D, 0x78)
BODY_COLOR = RGBColor(0x24, 0x2B, 0x36)
MUTED_COLOR = RGBColor(0x66, 0x72, 0x85)
TABLE_HEADER_FILL = "F2F4F7"
CODE_FILL = "F7F8FA"


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_element(element, fill: str) -> None:
    props = element.get_or_add_tcPr() if element.tag.endswith("tc") else element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    props.append(shd)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, HEADING_COLOR, 16, 8),
        ("Heading 2", 13, HEADING_COLOR, 12, 6),
        ("Heading 3", 12, HEADING_DARK_COLOR, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"Smart CRM 课程实训报告 | {title}")
    set_run_font(footer_run, size=9, color=MUTED_COLOR)


def add_title(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, size=22, bold=True, color=RGBColor(0x0B, 0x25, 0x45))

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("AI Sales Copilot CRM | 软件工程实训高分版交付文档")
    set_run_font(subtitle_run, size=10.5, color=MUTED_COLOR)


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    return text.strip()


def add_markdown_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(clean_inline_markdown(text))
    set_run_font(run, size=11, color=BODY_COLOR)


def add_list_item(doc: Document, text: str, ordered: bool = False) -> None:
    style = "List Number" if ordered else "List Bullet"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(clean_inline_markdown(text))
    set_run_font(run, size=11, color=BODY_COLOR)


def add_code_line(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    shade_element(paragraph._p, CODE_FILL)
    run = paragraph.add_run(text if text else " ")
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x3A, 0x45)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if not (stripped.startswith("|") and stripped.endswith("|")):
            continue
        cells = [clean_inline_markdown(cell) for cell in stripped.strip("|").split("|")]
        if all(re.fullmatch(r"\s*:?-{3,}:?\s*", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, table_lines: list[str]) -> None:
    rows = parse_table(table_lines)
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    for row in rows:
        row.extend([""] * (column_count - len(row)))

    table = doc.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    usable_width = 6.5
    first_col = 1.15 if column_count >= 3 else usable_width / column_count
    if column_count >= 3:
        remaining = usable_width - first_col
        widths = [first_col] + [remaining / (column_count - 1)] * (column_count - 1)
    else:
        widths = [usable_width / column_count] * column_count

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=9.2 if column_count >= 4 else 10, bold=row_index == 0, color=BODY_COLOR)
            if row_index == 0:
                shade_element(cell._tc, TABLE_HEADER_FILL)
    set_table_width(table, widths)
    doc.add_paragraph()


def markdown_to_docx(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    title = markdown_path.stem
    for line in lines:
        if line.startswith("# "):
            title = clean_inline_markdown(line[2:])
            break

    doc = Document()
    configure_document(doc, title)
    add_title(doc, title)

    in_code = False
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

    for line in lines:
        stripped = line.rstrip()
        if stripped.startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            add_code_line(doc, stripped)
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue
        flush_table()
        if not stripped:
            continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            doc.add_heading(clean_inline_markdown(stripped[3:]), level=1)
            continue
        if stripped.startswith("### "):
            doc.add_heading(clean_inline_markdown(stripped[4:]), level=2)
            continue
        if stripped.startswith("#### "):
            doc.add_heading(clean_inline_markdown(stripped[5:]), level=3)
            continue
        bullet_match = re.match(r"^\s*[-*]\s+(.+)$", stripped)
        if bullet_match:
            add_list_item(doc, bullet_match.group(1))
            continue
        number_match = re.match(r"^\s*\d+[.)]\s+(.+)$", stripped)
        if number_match:
            add_list_item(doc, number_match.group(1), ordered=True)
            continue
        add_markdown_paragraph(doc, stripped)

    flush_table()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def resolve_docs(args: argparse.Namespace) -> list[Path]:
    if args.docs:
        return [REPORT_ROOT / doc for doc in args.docs]
    if args.changed_only:
        return [
            REPORT_ROOT / "03_数据库设计文档.md",
            REPORT_ROOT / "04_后台接口设计文档.md",
            REPORT_ROOT / "05_软件实现说明.md",
            REPORT_ROOT / "07_迭代规划记录及项目总结.md",
            REPORT_ROOT / "08_黑盒白盒测试文档.md",
        ]
    return [REPORT_ROOT / doc for doc in DEFAULT_DOCS]


def main() -> None:
    parser = argparse.ArgumentParser(description="Build Smart CRM report DOCX files from Markdown drafts.")
    parser.add_argument("docs", nargs="*", help="Markdown filenames under 报告文档/v2-最终高分版")
    parser.add_argument("--changed-only", action="store_true", help="Build the documents touched by the latest owner-scope iteration.")
    args = parser.parse_args()

    for markdown_path in resolve_docs(args):
        if not markdown_path.exists():
            raise FileNotFoundError(markdown_path)
        output_name = OUTPUT_NAME_OVERRIDES.get(markdown_path.name, markdown_path.with_suffix(".docx").name)
        output_path = OUTPUT_ROOT / output_name
        markdown_to_docx(markdown_path, output_path)
        print(f"Built {output_path}")


if __name__ == "__main__":
    main()
